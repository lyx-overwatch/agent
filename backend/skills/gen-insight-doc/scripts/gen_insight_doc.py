#!/usr/bin/env python3
"""生成数字金融洞察格式 Word 文档"""
import sys
import re
import os
from copy import deepcopy
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ========== 格式常量 ==========
SKILL_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_FILE = os.path.join(SKILL_ROOT, 'templates', '数字金融洞察格式模板参考0413.docx')

FONT_TITLE_CN = '方正小标宋简体'
FONT_TITLE_EN = 'Times New Roman'
FONT_HEADING1_CN = '方正黑体'
FONT_HEADING2_CN = '方正楷体'
FONT_BODY_CN = '方正仿宋'
FONT_EN = 'Times New Roman'

SIZE_TITLE = Pt(22)
SIZE_BODY = Pt(16)
SIZE_TABLE = Pt(14)

COLOR_TITLE = RGBColor(0x2F, 0x78, 0xDA)
COLOR_BLACK = RGBColor(0, 0, 0)

INDENT_FIRST_LINE = Pt(SIZE_BODY.pt * 2)
LINE_SPACING = 1.2


def set_run_font(run, cn_font, en_font, size, bold=False, color=None):
    """设置 run 的字体属性"""
    run.font.name = en_font
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)


def add_paragraph(doc, text, cn_font, en_font, size, bold=False, color=None,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True,
                  line_spacing=LINE_SPACING):
    """添加一个格式化段落，支持 **加粗** 和 __加粗__ 内联标记"""
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_indent:
        pf.first_line_indent = INDENT_FIRST_LINE

    # 分割文本为 [('普通文本', False), ('加粗文本', True), ...]
    segments = re.split(r'(\*\*.+?\*\*|__.+?__)', text)
    for seg in segments:
        if not seg:
            continue
        is_bold = (seg.startswith('**') and seg.endswith('**')) or \
                  (seg.startswith('__') and seg.endswith('__'))
        if is_bold:
            seg = seg[2:-2]  # 去掉 **或__
            run = para.add_run(seg)
            set_run_font(run, cn_font, en_font, size, True, color or COLOR_BLACK)
        else:
            run = para.add_run(seg)
            set_run_font(run, cn_font, en_font, size, bold, color)
    return para


def update_header_textboxes(doc, issue_year=None, issue_number=None, issue_date=None):
    """更新模板头部文本框中的期号和日期"""
    import html as html_mod

    para = doc.paragraphs[0]
    xml_elem = para._element

    for t_elem in xml_elem.iter(qn('w:t')):
        text = t_elem.text
        if text is None:
            continue

        decoded = html_mod.unescape(text)

        if decoded.startswith('【') and '第' in decoded and '期' in decoded:
            if issue_year and issue_number:
                new_text = f'【{issue_year}】第{issue_number}期'
                t_elem.text = html_mod.escape(new_text) if '&' in new_text else new_text
                t_elem.set(qn('xml:space'), 'preserve')

        if re.match(r'\d{4}年\d+月\d+日', decoded):
            if issue_date:
                t_elem.text = html_mod.escape(issue_date) if '&' in issue_date else issue_date
                t_elem.set(qn('xml:space'), 'preserve')


def normalize_quotes(text):
    """将各类引号统一转换为中文引号，确保中文排版一致性。

    处理规则：
    1. ASCII 双引号 \" → 中文双引号 \"\"（成对替换）
    2. ASCII 单引号 ' → 中文单引号 ''（成对替换，跳过缩写如 don't）
    3. 英文智能引号 \"\" → 中文双引号 \"\"
    4. 低双引号 „ → 中文左双引号 \"
    5. 兜底扫描：处理中英文混排时可能残留的 ASCII 引号
    """
    # ── 第一轮：双引号 ——
    # Step 1.1: 英文智能双引号 / 低双引号 → 统一先转 ASCII
    text = text.replace('\u201c', '"').replace('\u201d', '"')  # " " → "
    text = text.replace('\u201e', '"')  # „ → "

    # Step 1.2: ASCII 双引号成对 → 中文双引号
    result = []
    in_dq = False
    for ch in text:
        if ch == '"':
            result.append('\u201c' if not in_dq else '\u201d')
            in_dq = not in_dq
        else:
            result.append(ch)
    text = ''.join(result)

    # ── 第二轮：单引号 ——
    # Step 2.1: 英文智能单引号 → 统一先转 ASCII
    text = text.replace('\u2018', "'").replace('\u2019', "'")  # ' ' → '

    # Step 2.2: ASCII 单引号成对 → 中文单引号（跳过缩写/所有格）
    result = []
    in_sq = False
    chars = list(text)
    for i, ch in enumerate(chars):
        if ch == "'":
            if in_sq:
                result.append('\u2019')  # 右单引号
                in_sq = False
            else:
                prev_char = chars[i - 1] if i > 0 else ''
                next_char = chars[i + 1] if i < len(chars) - 1 else ''
                if prev_char.isalpha() and next_char.isalpha():
                    # 缩写中的撇号（don't, it's, I'm 等），保留 ASCII
                    result.append("'")
                elif prev_char.isalpha() and not next_char.isalpha():
                    # 词尾所有格（owners', students' 等），保留 ASCII
                    result.append("'")
                elif prev_char.isdigit() and next_char.isdigit():
                    # 数字缩写（'90s, '80s 等），保留 ASCII
                    result.append("'")
                else:
                    result.append('\u2018')  # 左单引号
                    in_sq = True
        else:
            result.append(ch)
    text = ''.join(result)

    # ── 第三轮：兜底 ——
    # 残存的孤立 ASCII 双引号 → 中文左双引号（宁可错杀、不可漏网）
    leftover_dq = text.count('"')
    if leftover_dq > 0:
        text = text.replace('"', '\u201c')

    # 残存的孤立 ASCII 单引号（不在缩写位置）→ 中文左单引号
    if text.count("'") > 0:
        words = list(text)
        for i, ch in enumerate(words):
            if ch == "'":
                prev_alpha = i > 0 and words[i - 1].isalpha()
                next_alpha = i < len(words) - 1 and words[i + 1].isalpha()
                if not (prev_alpha and next_alpha):
                    words[i] = '\u2018'
        text = ''.join(words)

    return text


def parse_markdown(md_text):
    """解析 Markdown 文本，返回结构化内容列表"""
    md_text = normalize_quotes(md_text)
    elements = []
    lines = md_text.strip().split('\n')

    h1_pattern = re.compile(r'^[一二三四五六七八九十]+、')
    h2_pattern = re.compile(r'^（[一二三四五六七八九十]+）')
    h3_pattern = re.compile(r'^\d+\.\s')
    h4_pattern = re.compile(r'^（\d+）')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith('# ') and not stripped.startswith('## '):
            elements.append(('title', stripped[2:]))
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            elements.append(('h1', stripped[3:]))
        elif stripped.startswith('### ') and not stripped.startswith('#### '):
            elements.append(('h2', stripped[4:]))
        elif stripped.startswith('#### '):
            elements.append(('h3', stripped[5:]))
        elif h1_pattern.match(stripped):
            elements.append(('h1', stripped))
        elif h2_pattern.match(stripped):
            elements.append(('h2', stripped))
        elif h4_pattern.match(stripped):
            elements.append(('h4', stripped))
        elif h3_pattern.match(stripped):
            elements.append(('h3', stripped))
        elif stripped.startswith('内容提供'):
            elements.append(('provider', stripped))
        else:
            elements.append(('body', stripped))

    return elements


def generate_docx(md_text, output_path, title=None, provider=None,
                  issue_year=None, issue_number=None, issue_date=None):
    """从 Markdown 文本生成格式化 Word 文档"""

    template_path = TEMPLATE_FILE
    if os.path.exists(template_path):
        doc = Document(template_path)
        update_header_textboxes(doc, issue_year, issue_number, issue_date)
        body = doc.element.body
        paras_to_remove = []
        p_idx = 0
        for child in list(body):
            if child.tag == qn('w:p'):
                if p_idx < 4:
                    p_idx += 1
                    continue
                paras_to_remove.append(child)
                p_idx += 1
        for p in paras_to_remove:
            body.remove(p)
    else:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    elements = parse_markdown(md_text)

    if title:
        elements = [(t, c) if t != 'title' else ('title', title) for t, c in elements]
        if not any(t == 'title' for t, c in elements):
            elements.insert(0, ('title', title))
    if provider:
        elements = [(t, c) if t != 'provider' else ('provider', provider) for t, c in elements]

    for elem_type, text in elements:
        if elem_type == 'title':
            add_paragraph(doc, text, FONT_TITLE_CN, FONT_TITLE_EN, SIZE_TITLE,
                          bold=True, color=COLOR_TITLE,
                          align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
        elif elem_type == 'provider':
            add_paragraph(doc, text, FONT_BODY_CN, FONT_EN, SIZE_BODY,
                          bold=True, color=COLOR_BLACK,
                          align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
        elif elem_type == 'h1':
            add_paragraph(doc, text, FONT_HEADING1_CN, FONT_EN, SIZE_BODY,
                          bold=False, color=COLOR_BLACK, first_indent=True)
        elif elem_type == 'h2':
            add_paragraph(doc, text, FONT_HEADING2_CN, FONT_EN, SIZE_BODY,
                          bold=False, color=COLOR_BLACK, first_indent=True)
        elif elem_type == 'h3':
            add_paragraph(doc, text, FONT_BODY_CN, FONT_EN, SIZE_BODY,
                          bold=False, color=COLOR_BLACK, first_indent=True)
        elif elem_type == 'h4':
            add_paragraph(doc, text, FONT_BODY_CN, FONT_EN, SIZE_BODY,
                          bold=False, color=COLOR_BLACK, first_indent=True)
        else:
            add_paragraph(doc, text, FONT_BODY_CN, FONT_EN, SIZE_BODY,
                          bold=False, color=COLOR_BLACK, first_indent=True)

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    input_file = sys.argv[1] if len(sys.argv) > 1 else 'insight_report.md'
    output_file = sys.argv[2] if len(sys.argv) > 2 else '数字金融洞察_政绩观与数字金融.docx'
    issue_year = sys.argv[3] if len(sys.argv) > 3 else '2026'
    issue_number = sys.argv[4] if len(sys.argv) > 4 else '5'
    issue_date = sys.argv[5] if len(sys.argv) > 5 else '2026年5月11日'

    with open(input_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 写入前清洗源文件：引号规范化
    cleaned = normalize_quotes(md_text)
    if cleaned != md_text:
        with open(input_file, 'w', encoding='utf-8') as f:
            f.write(cleaned)
        print(f"🔤 已清洗源文件引号: {input_file}")

    generate_docx(cleaned, output_file, issue_year=issue_year,
                  issue_number=issue_number, issue_date=issue_date)
    print(f"文档已生成: {output_file}")
