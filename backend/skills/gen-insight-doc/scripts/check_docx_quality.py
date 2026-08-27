#!/usr/bin/env python3
"""检查并修复生成的 docx 格式质量：字体合规 + 异常字符

用法:
    python3 check_docx_quality.py <output.docx>          # 只检查
    python3 check_docx_quality.py <output.docx> --fix    # 检查 + 修复
"""
import re
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

EXPECTED_CN_FONTS = {'方正小标宋简体', '方正黑体', '方正楷体', '方正仿宋', '华文行楷'}
EXPECTED_EN_FONTS = {'Times New Roman'}

# ========== 段落类型 → 字体规范 ==========
FONT_BY_TYPE = {
    'title': {
        'cn': '方正小标宋简体', 'en': 'Times New Roman', 'size': Pt(22),
        'bold': True, 'color': RGBColor(0x2F, 0x78, 0xDA),
    },
    'content_provider': {
        'cn': '方正仿宋', 'en': 'Times New Roman', 'size': Pt(16), 'bold': True,
    },
    'heading1': {
        'cn': '方正黑体', 'en': 'Times New Roman', 'size': Pt(16),
    },
    'heading2': {
        'cn': '方正楷体', 'en': 'Times New Roman', 'size': Pt(16),
    },
    'body': {
        'cn': '方正仿宋', 'en': 'Times New Roman', 'size': Pt(16),
    },
}

HEADING1_PREFIX = tuple(f'{i}、' for i in
    '一二三四五六七八九十')
HEADING2_PREFIX = tuple(f'（{i}）' for i in
    '一二三四五六七八九十')

def detect_para_type(para):
    """根据段落内容和格式推断段落类型"""
    text = para.text.strip()
    if not text:
        return 'body'

    # 文档大标题: 蓝色文字
    for run in para.runs:
        if run.font.color and run.font.color.rgb:
            if run.font.color.rgb == RGBColor(0x2F, 0x78, 0xDA):
                return 'title'

    # 内容提供行
    if text.startswith('内容提供'):
        return 'content_provider'

    # 一级标题: 一、二、三、...
    if text.startswith(HEADING1_PREFIX):
        return 'heading1'

    # 二级标题: （一）（二）（三）...
    if text.startswith(HEADING2_PREFIX):
        return 'heading2'

    return 'body'


# ========== 字体修复 ==========

def fix_run_fonts(run, cn_font, en_font, size=None, bold=None, color=None):
    """将 run 的字体修正为指定值"""
    run.font.name = en_font
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color

    # XML 层面设置中文字体
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


# ========== 字符清理 ==========
FIX_RULES = [
    (r'\*{3,}', '', '连续星号 ***+'),
    (r'#{3,}', '', '连续井号 ###+'),
    (r'_{3,}', '', '连续下划线 ___+'),
    (r'~{3,}', '', '连续波浪线 ~~~+'),
    (r'\*\*(.+?)\*\*', r'\1', 'Markdown 加粗 **text**'),
    (r'__(.+?)__', r'\1', 'Markdown 加粗 __text__'),
    (r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', 'Markdown 斜体 *text*'),
    (r'(?<!_)_(?!_)(.+?)(?<!_)_(?!_)', r'\1', 'Markdown 斜体 _text_'),
    (r'[\u200b\u200c\u200d\u200e\u200f]+', '', '零宽字符'),
    (r'\ufffd+', '', '乱码字符 �'),
]
FIX_BOLD_PATTERNS = [r'\*\*(.+?)\*\*', r'__(.+?)__']


# ========== 引号规范化 ==========

def normalize_quotes(text):
    """将各类引号统一转换为中文引号（与 gen_insight_doc.py 同逻辑）"""
    # 英文智能双引号 / 低双引号 → ASCII
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u201e', '"')

    # ASCII 双引号成对 → 中文双引号
    result = []
    in_dq = False
    for ch in text:
        if ch == '"':
            result.append('\u201c' if not in_dq else '\u201d')
            in_dq = not in_dq
        else:
            result.append(ch)
    text = ''.join(result)

    # 英文智能单引号 → ASCII
    text = text.replace('\u2018', "'").replace('\u2019', "'")

    # ASCII 单引号成对 → 中文单引号（跳过缩写）
    result = []
    in_sq = False
    chars = list(text)
    for i, ch in enumerate(chars):
        if ch == "'":
            if in_sq:
                result.append('\u2019')
                in_sq = False
            else:
                prev_c = chars[i - 1] if i > 0 else ''
                next_c = chars[i + 1] if i < len(chars) - 1 else ''
                if prev_c.isalpha() and next_c.isalpha():
                    result.append("'")
                elif prev_c.isalpha() and not next_c.isalpha():
                    result.append("'")
                elif prev_c.isdigit() and next_c.isdigit():
                    result.append("'")
                else:
                    result.append('\u2018')
                    in_sq = True
        else:
            result.append(ch)
    text = ''.join(result)

    # 兜底：残存的孤立双引号 → 左双引号
    if '"' in text:
        text = text.replace('"', '\u201c')
    # 兜底：非缩写的单引号 → 左单引号
    if "'" in text:
        words = list(text)
        for i, ch in enumerate(words):
            if ch == "'":
                p = words[i - 1] if i > 0 else ''
                n = words[i + 1] if i < len(words) - 1 else ''
                if not (p.isalpha() and n.isalpha()):
                    words[i] = '\u2018'
        text = ''.join(words)

    return text


QUOTE_CHECK_PATTERNS = [
    (r'\u0022', 'ASCII 双引号 \"', True),
    (r"(?<!\w)'(?!\w)", 'ASCII 单引号用作引号', True),
]


# ========== 检查 + 修复 ==========

def process_docx(doc, fix_mode=False):
    """扫描并（可选）修复 docx，返回 (issues, warnings, fonts_fixed, chars_fixed)"""
    issues = []
    warnings = []
    fonts_fixed = 0
    chars_fixed = 0

    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        p_label = f"第{p_idx+1}段" if p_idx < 50 else f"P{p_idx+1}"
        para_type = detect_para_type(para) if fix_mode else None
        font_spec = FONT_BY_TYPE.get(para_type, FONT_BY_TYPE['body']) if para_type else None

        for r_idx, run in enumerate(para.runs):
            if not run.text.strip():
                continue

            cn_font, en_font = get_run_fonts(run)
            run_label = f"{p_label} run{r_idx+1}"

            # ---- 字体检查 / 修复 ----
            if fix_mode and font_spec:
                # 修复：直接设为正确字体
                needs_fix = (
                    cn_font != font_spec['cn'] or
                    en_font != font_spec['en']
                )
                if needs_fix:
                    fix_run_fonts(
                        run,
                        cn_font=font_spec['cn'],
                        en_font=font_spec['en'],
                        size=font_spec.get('size'),
                        bold=font_spec.get('bold'),
                        color=font_spec.get('color'),
                    )
                    fonts_fixed += 1
            else:
                # 只检查不修复
                if cn_font is None:
                    issues.append(f"❌ {run_label}: 未设置中文字体 → 内容：\"{run.text[:30]}...\"")
                elif cn_font not in EXPECTED_CN_FONTS:
                    # 尝试推断应该用什么字体
                    para_type = detect_para_type(para)
                    expected_cn = FONT_BY_TYPE[para_type]['cn']
                    issues.append(f"❌ {run_label}: 中文字体 \"{cn_font}\"，应为 \"{expected_cn}\"（{para_type}）。内容：\"{run.text[:30]}...\"")

                if en_font is not None and en_font not in EXPECTED_EN_FONTS:
                    warnings.append(f"⚠️  {run_label}: 西文字体 \"{en_font}\"，预期 Times New Roman")

            # ---- 字符检查 / 修复 ----
            if fix_mode:
                for pattern, replacement, desc in FIX_RULES:
                    if re.search(pattern, run.text):
                        old_text = run.text
                        run.text = re.sub(pattern, replacement, run.text)
                        if run.text != old_text:
                            chars_fixed += 1
                            for bp in FIX_BOLD_PATTERNS:
                                if re.match(f'.*{bp}.*', old_text):
                                    run.font.bold = True
                                    break

                # ════ 引号修复 ════
                old_text = run.text
                run.text = normalize_quotes(run.text)
                if run.text != old_text:
                    chars_fixed += 1
            else:
                for pattern, replacement, desc in FIX_RULES:
                    for m in re.finditer(pattern, run.text):
                        issues.append(f"❌ {run_label}: 发现 {desc} \"{m.group()}\"")

                # ════ 引号检查 ════
                for pattern, desc, _ in QUOTE_CHECK_PATTERNS:
                    matches = list(re.finditer(pattern, run.text))
                    for m in matches:
                        issues.append(f"❌ {run_label}: 发现 {desc} \"{m.group()}\"，应为中文引号")

    return issues, warnings, fonts_fixed, chars_fixed


def get_run_fonts(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None, None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return None, None
    cn = rFonts.get(qn('w:eastAsia'))
    en = rFonts.get(qn('w:ascii'))
    return cn, en


# ========== 入口 ==========

def main():
    fix_mode = '--fix' in sys.argv
    args = [a for a in sys.argv[1:] if a != '--fix']

    if len(args) < 1:
        print("用法: python3 check_docx_quality.py <output.docx> [--fix]")
        sys.exit(1)

    filepath = args[0]
    doc = Document(filepath)

    if fix_mode:
        print(f"🔧 修复模式：{filepath}\n")
        issues, warnings, fonts_fixed, chars_fixed = process_docx(doc, fix_mode=True)
        doc.save(filepath)

        print(f"✅ 已保存修复后的文件: {filepath}")
        if fonts_fixed:
            print(f"🔤 字体修正: {fonts_fixed} 处")
        if chars_fixed:
            print(f"📝 字符清理: {chars_fixed} 处")
        if not fonts_fixed and not chars_fixed:
            print("  无需修复")
    else:
        print(f"📄 检查文件: {filepath}\n")
        issues, warnings, _, _ = process_docx(doc, fix_mode=False)

    print(f"\n📊 共 {len(doc.paragraphs)} 个段落\n")

    if issues:
        print(f"🔴 问题 {len(issues)} 处（需人工处理）：")
        for i, issue in enumerate(issues, 1):
            print(f"  {i}. {issue}")
        print()

    if warnings:
        print(f"🟡 提示 {len(warnings)} 处：")
        for i, w in enumerate(warnings, 1):
            print(f"  {i}. {w}")
        print()

    if not issues and not warnings:
        print("✅ 格式检查全部通过！")

    return len(issues)


if __name__ == '__main__':
    n = main()
    sys.exit(0 if n == 0 else 1)
